import os
import docx
from project.documents.section import Section
from configparser import ConfigParser, ExtendedInterpolation

# configuration
config = ConfigParser(interpolation=ExtendedInterpolation())
config.read('../config.ini')
IN_PROGRESS_PATH = config['DOCX']['IN_PROGRESS_PATH']
COMPLETED_PATH = config['DOCX']['COMPLETED_PATH']


class Document:
    """ extract and structure the text from a single docx file into the following:

     doc_text: single string with the full text of the document
     sections: document text structured in sections using heuristics from section_criteria module
     tables: text from the tables in a document
     properties: document metadata such as author_name and created_date

     NOTE: only compatible with docx, not earlier Microsoft Word file formats like doc
     """

    def __init__(self, path, doc_text=True, sections=True, table_text=True, doc_properties=True):
        """
        :param path: directory path to document
        :param doc_text: if True, extract document text
        :param sections: if True, structure document text into sections
        :param doc_properties: if True, extract all document properties (e.g. author_name)
        :param table_text: if True, extract the text from document tables
        """

        self.filename = path
        self.paragraphs = None
        self.text = None
        self.table_text = None
        self.author = None
        self.last_modified_by = None
        self.created = None
        self.last_printed = None
        self.revision = None
        self.num_tables = None
        self.path = os.path.join(IN_PROGRESS_PATH, path)
        self.completed_path = os.path.join(COMPLETED_PATH, path)

        self.sections = []  # container for section objects

        # use docx to read document xml
        doc_current_path = os.path.join(IN_PROGRESS_PATH, path)
        doc = docx.Document(doc_current_path)

        if doc_text:
            self.paragraphs = doc.paragraphs
            self.text = ' '.join([para.text.strip() for para in self.paragraphs if para.text.strip() != ''])

        if sections:
            self._set_sections(doc)

        if table_text:
            self._set_table_text(doc)

        if doc_properties:
            self._set_doc_properties(doc)

    def __repr__(self):
        return "<DOCUMENT: {}>".format(self.filename)

    def get_sections_dict(self):
        """ extract the section_name: section_text key value pairs from the sections objects """
        sections_dict = {}
        for section in self.sections:
            sections_dict[section.section_name] = {'criteria':section.criteria, 'text': section.get_section_text()}

        return sections_dict

    def _set_sections(self, doc):
        """ structure the document text into sections

        :param doc: docx document
        """

        # set a name for the first section to store all the text in case
        # none of the section criteria are met in the document
        section = Section(section_name='First Section')

        for p in doc.paragraphs:
            if section.paragraph_doesnt_have_text(p, alpha_only=True):
                continue
            if section.is_section_header(p):
                if not section.section_has_text():
                    # section text is still empty so
                    # the section name is spread across multiple lines
                    section.set_section_name(p.text)
                else:
                    # section text is fully populated so
                    # reset for the next section
                    self.sections.append(section)
                    section = Section(section_name=p.text)
            else:
                # not a new section, so continue collecting text for the current section
                section.add_section_text(p.text)

        # add final section - necessary as sections are only appended to
        # the sections list after the next section header is identified
        if section.section_has_text():
            self.sections.append(section)

    def _set_table_text(self, doc):
        """ set the table_text attribute with text from the documents tables

        :param doc: docx document
        """

        table_text = []
        for table in doc.tables:
            try:
                for cell in table._cells:
                   table_text.append(cell.text.strip())
            except IndexError:
                continue

        self.table_text = ' '.join(table_text)

    def _set_doc_properties(self, doc):
        self.author = doc.core_properties.author
        self.last_modified_by = doc.core_properties.last_modified_by
        self.created = doc.core_properties.created
        self.last_printed = doc.core_properties.last_printed
        self.revision = doc.core_properties.revision
        self.num_tables = len(doc.tables)
